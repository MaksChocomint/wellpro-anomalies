from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
REPO_ROOT = ROOT.parents[1]
SRC_MD = ROOT / "wellpro_vkr_bachelor_draft_20260427.md"
OUT_DOCX = ROOT / "wellpro_vkr_mordvinova_edited_20260427.docx"
THESIS_FIGURES = ROOT / "thesis_figures"
EXTRACTED_ASSETS = ROOT / "extracted_assets"


START_MARK = "## 2. Черновик полного текста ВКР"
END_MARK = "## 3. Список мест для рисунков и схем"


CHAPTER_PAGE_BREAKS = {
    "## 1 Аналитический раздел",
    "## 2 Проектный раздел",
    "## 3 Экспериментальный раздел",
    "## 4 Экономический раздел",
    "## 5 Безопасность и экологичность",
    "## Заключение",
    "## Список использованных источников",
}


FONT_MAIN = "Times New Roman"
FONT_TITLE = "Garamond"


numbered_re = re.compile(r"^(\d+)\.\s+(.*)$")
figure_placeholder_re = re.compile(r"^\[Место для рисунка\s+([\d\.]+)\s+-\s+(.+)\]$")
table_caption_re = re.compile(r"^Таблица\s+[\d\.]+\s+-\s+.+$")


FIGURE_ASSETS = {
    "2.1": {
        "path": EXTRACTED_ASSETS / "tz_01_software_structure.png",
        "width_cm": 16.0,
        "source": "Источник - составлено по материалам «ТЗ 06.08.2025.docx».",
    },
    "2.2": {
        "path": EXTRACTED_ASSETS / "tz_02_app_workflow.png",
        "width_cm": 16.0,
        "source": "Источник - составлено по материалам «ТЗ 06.08.2025.docx».",
    },
    "2.3": {
        "path": THESIS_FIGURES / "figure_02_db_models.png",
        "width_cm": 16.0,
        "source": "Источник - составлено по материалам «Курсовая_работа_БД_ИАС_22_Иванов_М_С.docx».",
    },
    "2.4": {
        "path": REPO_ROOT / "frontend" / "qa_after_company.png",
        "width_cm": 16.0,
        "source": "Источник - скриншот текущей версии приложения из репозитория проекта.",
    },
    "2.5": {
        "path": THESIS_FIGURES / "figure_04_monitoring_interface.png",
        "width_cm": 16.0,
        "source": "Источник - составлено по реальным скриншотам интерфейса из курсового проекта по БД.",
    },
    "3.1": {
        "path": THESIS_FIGURES / "figure_05_ammad_example.png",
        "width_cm": 16.0,
        "source": "Источник - построено автором по файлу «ammad_case_03_context_stuck_sensor.TXT».",
    },
    "3.2": {
        "path": THESIS_FIGURES / "figure_06_method_comparison.png",
        "width_cm": 16.0,
        "source": "Источник - построено автором по benchmark на 15 файлах и validation cases.",
    },
}


TITLE_LAYOUT = {
    "top": Cm(1.0),
    "bottom": Cm(2.5),
    "left": Cm(2.5),
    "right": Cm(1.0),
    "header": Cm(1.27),
    "footer": Cm(1.27),
}

BODY_LAYOUT = {
    "top": Cm(2.0),
    "bottom": Cm(2.0),
    "left": Cm(3.0),
    "right": Cm(1.5),
    "header": Cm(1.25),
    "footer": Cm(1.25),
}


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def apply_section_layout(section, layout: dict[str, Cm]) -> None:
    section.top_margin = layout["top"]
    section.bottom_margin = layout["bottom"]
    section.left_margin = layout["left"]
    section.right_margin = layout["right"]
    section.header_distance = layout["header"]
    section.footer_distance = layout["footer"]


def set_run_font(
    run,
    *,
    size: int = 12,
    bold: bool | None = None,
    italic: bool | None = None,
    font_name: str = FONT_MAIN,
) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_paragraph(
    paragraph,
    *,
    align: WD_PARAGRAPH_ALIGNMENT | None = None,
    first_line_indent_cm: float = 1.25,
    line_spacing: float = 1.5,
    before_pt: int = 0,
    after_pt: int = 0,
) -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.space_after = Pt(after_pt)
    paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)


def add_text_paragraph(
    doc: Document,
    text: str,
    *,
    size: int = 12,
    bold: bool = False,
    italic: bool = False,
    align: WD_PARAGRAPH_ALIGNMENT | None = WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    first_line_indent_cm: float = 1.25,
    line_spacing: float = 1.5,
    before_pt: int = 0,
    after_pt: int = 0,
    font_name: str = FONT_MAIN,
    style: str | None = None,
) -> None:
    p = doc.add_paragraph(style=style)
    configure_paragraph(
        p,
        align=align,
        first_line_indent_cm=first_line_indent_cm,
        line_spacing=line_spacing,
        before_pt=before_pt,
        after_pt=after_pt,
    )
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, font_name=font_name)


def add_blank_paragraph(doc: Document, *, after_pt: int = 0) -> None:
    p = doc.add_paragraph()
    configure_paragraph(
        p,
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        first_line_indent_cm=0,
        line_spacing=1.0,
        after_pt=after_pt,
    )


def add_image_paragraph(doc: Document, image_path: Path, *, width_cm: float = 16.0) -> None:
    p = doc.add_paragraph()
    configure_paragraph(
        p,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def add_source_note(doc: Document, text: str) -> None:
    add_text_paragraph(
        doc,
        text,
        size=10,
        italic=True,
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )


def style_base_document(doc: Document) -> None:
    clear_document(doc)
    apply_section_layout(doc.sections[0], TITLE_LAYOUT)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_MAIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MAIN)
    normal.font.size = Pt(12)

    heading_1 = doc.styles["Heading 1"]
    heading_1.font.name = FONT_MAIN
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MAIN)
    heading_1.font.size = Pt(14)
    heading_1.font.bold = True

    for style_name in ("Heading 2", "Heading 3", "Heading 4"):
        style = doc.styles[style_name]
        style.font.name = FONT_MAIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MAIN)
        style.font.size = Pt(12)
        style.font.bold = True

    for style_name in ("List Bullet", "List Number", "toc 1", "toc 2", "toc 3"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = FONT_MAIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MAIN)
        style.font.size = Pt(12)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_new_section(doc: Document, layout: dict[str, Cm]) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_section_layout(section, layout)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = " "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(begin)
    r.append(instr)
    r.append(separate)
    r.append(placeholder)
    r.append(end)
    set_run_font(run)


def add_body_heading(doc: Document, text: str, level: int) -> None:
    style_name = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(style=style_name)
    align = WD_PARAGRAPH_ALIGNMENT.CENTER if level == 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
    size = 14 if level == 1 else 12
    configure_paragraph(
        p,
        align=align,
        first_line_indent_cm=0,
        line_spacing=1.5,
        before_pt=0,
        after_pt=0,
    )
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)


def add_front_heading(doc: Document, text: str) -> None:
    add_text_paragraph(
        doc,
        text,
        size=14,
        bold=False,
        italic=False,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.5,
    )


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    configure_paragraph(
        p,
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        first_line_indent_cm=0,
        line_spacing=1.5,
    )
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')


def add_footer_page_numbers(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_field(p, "PAGE")


def extract_body(md_text: str) -> list[str]:
    start = md_text.index(START_MARK) + len(START_MARK)
    end = md_text.index(END_MARK)
    return md_text[start:end].strip().splitlines()


def collect_until_next_heading(lines: list[str], start: int) -> tuple[list[str], int]:
    block: list[str] = []
    i = start
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith("### ") or stripped.startswith("## "):
            break
        block.append(lines[i])
        i += 1
    return block, i


def clean_md_line(text: str) -> str:
    return text.strip().replace("  ", " ").strip("*").strip()


def parse_title_info(block: list[str]) -> dict[str, str]:
    lines = [clean_md_line(line) for line in block if clean_md_line(line)]

    def find(prefix: str, default: str = "") -> str:
        for line in lines:
            if line.startswith(prefix):
                return line.split(":", 1)[1].strip()
        return default

    topic = "«Разработка информационной системы детекции и анализа аномальных показаний сенсорных датчиков буровых установок»"
    if "Тема:" in lines:
        idx = lines.index("Тема:")
        if idx + 1 < len(lines):
            topic = lines[idx + 1]

    return {
        "direction": find("Направление подготовки:", "09.03.01 - Информатика и вычислительная техника"),
        "profile": find("Профиль:", "____________________________________________"),
        "topic": topic,
        "author": find("Автор:", "____________________________________________"),
        "group": find("Группа:", "____________________________________________"),
        "supervisor": find("Руководитель:", "____________________________________________"),
        "city": lines[-2] if len(lines) >= 2 else "Санкт-Петербург",
        "year": lines[-1] if lines else "2026",
    }


def render_title_page(doc: Document, info: dict[str, str]) -> None:
    add_text_paragraph(
        doc,
        "Министерство науки и высшего образования Российской Федерации",
        size=12,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "Санкт-Петербургский горный университет императрицы Екатерины II",
        size=14,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_blank_paragraph(doc)
    add_text_paragraph(
        doc,
        "Кафедра Информационных систем и вычислительной техники",
        size=12,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_blank_paragraph(doc, after_pt=6)
    add_text_paragraph(
        doc,
        "Допускается к защите в ГЭК",
        size=14,
        bold=True,
        italic=True,
        align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "Зав. кафедрой ИСиВТ",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "____________________________________________",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "«____» ______________ 2026 г.",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_blank_paragraph(doc, after_pt=8)
    add_text_paragraph(
        doc,
        "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА",
        size=18,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
        font_name=FONT_TITLE,
    )
    add_text_paragraph(
        doc,
        "(дипломная работа бакалавра)",
        size=18,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
        font_name=FONT_TITLE,
    )
    add_blank_paragraph(doc, after_pt=8)
    add_text_paragraph(
        doc,
        f"на тему: {info['topic']}",
        size=14,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "Утверждена Приказом от __________  № _________",
        size=12,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
        before_pt=6,
    )
    add_blank_paragraph(doc)
    add_text_paragraph(
        doc,
        info["direction"],
        size=12,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "(шифр)                         (наименование направления)",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        info["profile"],
        size=12,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "(наименование направленности (профиля))",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_blank_paragraph(doc)
    add_text_paragraph(
        doc,
        f"{info['author']}    ____________________",
        size=14,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "(Ф.И.О.)                         (подпись)",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        f"{info['supervisor']}    ____________________",
        size=14,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "(должность, звание)          (подпись)",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "Рецензент:  _________________________________________________",
        size=14,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "(должность, звание)          (подпись)",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_blank_paragraph(doc, after_pt=10)
    add_text_paragraph(
        doc,
        info["city"],
        size=14,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        info["year"],
        size=14,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )


def render_agreement_page(doc: Document, info: dict[str, str]) -> None:
    add_text_paragraph(
        doc,
        "ЛИСТ СОГЛАСОВАНИЯ",
        size=14,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "по разделам выпускной квалификационной работы",
        size=14,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.0,
        before_pt=14,
        after_pt=14,
    )
    add_text_paragraph(
        doc,
        f"на тему: {info['topic']}",
        size=14,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_blank_paragraph(doc)
    add_text_paragraph(
        doc,
        "Автор: студент гр. _______________    ____________________",
        size=14,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "(шифр)                         (подпись)",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent_cm=0,
        line_spacing=1.0,
        after_pt=6,
    )

    section_notes = [
        "связанного с экономическим обоснованием *",
        "связанного с вопросами экологии *",
        "связанного с вопросами охраны труда и безопасностью жизнедеятельности *",
        "связанного с вопросами качества перевода на иностранный язык *",
    ]

    for note in section_notes:
        add_text_paragraph(
            doc,
            "Раздел _______________________________________________________________",
            size=12,
            align=WD_PARAGRAPH_ALIGNMENT.LEFT,
            first_line_indent_cm=0,
            line_spacing=1.0,
        )
        add_text_paragraph(
            doc,
            f"(наименование раздела, {note})",
            size=12,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            first_line_indent_cm=0,
            line_spacing=1.0,
        )
        add_text_paragraph(
            doc,
            "Кафедра__________               _______________     ________________________________",
            size=12,
            align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent_cm=0,
            line_spacing=1.0,
        )
        add_text_paragraph(
            doc,
            "(название кафедры)          (должность, звание)              (подпись)",
            size=12,
            align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            first_line_indent_cm=0,
            line_spacing=1.0,
            after_pt=6,
        )

    add_blank_paragraph(doc, after_pt=6)
    add_text_paragraph(
        doc,
        "* Разделы ВКР и их наименование определяются в соответствии со структурой и направленностью выполняемой работы.",
        size=10,
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        first_line_indent_cm=0,
        line_spacing=1.0,
    )


def add_picture_placeholder(doc: Document, number: str, caption: str) -> None:
    asset = FIGURE_ASSETS.get(number)
    if asset and Path(asset["path"]).exists():
        add_image_paragraph(doc, Path(asset["path"]), width_cm=float(asset.get("width_cm", 16.0)))
    else:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.rows[0].height = Cm(5.5)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        cell_p = cell.paragraphs[0]
        configure_paragraph(
            cell_p,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            first_line_indent_cm=0,
            line_spacing=1.0,
        )
        set_run_font(cell_p.add_run(f"Место для рисунка {number}"), size=12)
    add_text_paragraph(
        doc,
        f"Рисунок {number} - {caption}",
        size=12,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_cm=0,
        line_spacing=1.5,
    )
    if asset and asset.get("source"):
        add_source_note(doc, str(asset["source"]))


def build_docx(lines: list[str]) -> Document:
    doc = Document()
    style_base_document(doc)

    title_info: dict[str, str] | None = None
    current_section = ""
    i = 0
    body_started = False

    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            i += 1
            continue

        if stripped == "### Титульный лист":
            block, i = collect_until_next_heading(lines, i + 1)
            title_info = parse_title_info(block)
            render_title_page(doc, title_info)
            continue

        if stripped == "### Лист согласования":
            if title_info is None:
                title_info = parse_title_info([])
            add_new_section(doc, TITLE_LAYOUT)
            render_agreement_page(doc, title_info)
            block, i = collect_until_next_heading(lines, i + 1)
            continue

        if stripped == "### Аннотация":
            add_new_section(doc, BODY_LAYOUT)
            add_front_heading(doc, "Аннотация")
            current_section = "annotation"
            i += 1
            continue

        if stripped == "### Abstract":
            add_page_break(doc)
            add_front_heading(doc, "Abstract")
            current_section = "abstract"
            i += 1
            continue

        if stripped in {"### Содержание", "### Оглавление"}:
            add_page_break(doc)
            add_front_heading(doc, "ОГЛАВЛЕНИЕ")
            add_toc(doc)
            current_section = "toc"
            block, i = collect_until_next_heading(lines, i + 1)
            continue

        if stripped == "### Введение":
            add_new_section(doc, BODY_LAYOUT)
            add_footer_page_numbers(doc.sections[-1])
            add_body_heading(doc, "Введение", 1)
            current_section = "body"
            body_started = True
            i += 1
            continue

        if stripped.startswith("## "):
            if stripped in CHAPTER_PAGE_BREAKS and body_started:
                add_page_break(doc)
            add_body_heading(doc, stripped[3:], 1)
            current_section = "body"
            body_started = True
            i += 1
            continue

        if stripped.startswith("### "):
            add_body_heading(doc, stripped[4:], 2)
            current_section = "body"
            body_started = True
            i += 1
            continue

        if stripped.startswith("#### "):
            add_body_heading(doc, stripped[5:], 3)
            current_section = "body"
            body_started = True
            i += 1
            continue

        figure_match = figure_placeholder_re.match(stripped)
        if figure_match:
            number, caption = figure_match.groups()
            add_picture_placeholder(doc, number, caption)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
                headers = rows[0]
                data_rows = rows[2:]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for idx, header in enumerate(headers):
                    table.rows[0].cells[idx].text = header
                for row in data_rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row):
                        cells[idx].text = value
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            configure_paragraph(
                                paragraph,
                                align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                                first_line_indent_cm=0,
                                line_spacing=1.0,
                            )
                            for run in paragraph.runs:
                                set_run_font(run, size=12)
                add_blank_paragraph(doc)
                continue

            for row in table_lines:
                add_text_paragraph(doc, row)
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            configure_paragraph(
                p,
                align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                first_line_indent_cm=0,
                line_spacing=1.5,
            )
            set_run_font(p.add_run(stripped[2:]), size=12)
            i += 1
            continue

        match = numbered_re.match(stripped)
        if match:
            p = doc.add_paragraph(style="List Number")
            configure_paragraph(
                p,
                align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                first_line_indent_cm=0,
                line_spacing=1.5,
            )
            set_run_font(p.add_run(match.group(2)), size=12)
            i += 1
            continue

        if table_caption_re.match(stripped):
            add_text_paragraph(
                doc,
                stripped,
                size=12,
                align=WD_PARAGRAPH_ALIGNMENT.LEFT,
                first_line_indent_cm=0,
            )
            i += 1
            continue

        if stripped.startswith("Рисунок "):
            add_text_paragraph(
                doc,
                stripped,
                size=12,
                align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                first_line_indent_cm=0,
            )
            i += 1
            continue

        if current_section == "toc":
            i += 1
            continue

        add_text_paragraph(doc, stripped)
        i += 1

    return doc


def main() -> None:
    md_text = SRC_MD.read_text(encoding="utf-8")
    lines = extract_body(md_text)
    doc = build_docx(lines)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
