from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
FILES = [
    ("wellpro_user_guide_20260427.md", "wellpro_user_guide_20260427.docx"),
    ("wellpro_admin_guide_20260427.md", "wellpro_admin_guide_20260427.docx"),
    ("wellpro_appendices_package_20260427.md", "wellpro_appendices_package_20260427.docx"),
]


numbered_re = re.compile(r"^(\d+)\.\s+(.*)$")
table_caption_re = re.compile(r"^Таблица\s+[\d\.]+\s+-\s+.+$")


def set_run_font(run, size: int = 14, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(56.7)
    section.bottom_margin = Pt(56.7)
    section.left_margin = Pt(85.05)
    section.right_margin = Pt(42.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True


def add_paragraph(doc: Document, text: str, *, style: str = "Normal", first_indent: bool = True) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(35.4) if first_indent and style == "Normal" else Pt(0)
    run = p.add_run(text)
    set_run_font(run)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=True)


def build_doc(src: Path, dst: Path) -> None:
    text = src.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    setup_doc(doc)

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 1)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 2)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 3)
            i += 1
            continue
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:], 4)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
                headers = rows[0]
                data_rows = rows[2:]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for idx, header in enumerate(headers):
                    table.rows[0].cells[idx].text = header
                for row in data_rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row):
                        cells[idx].text = value
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.line_spacing = 1.0
                            for run in paragraph.runs:
                                set_run_font(run, size=12)
                doc.add_paragraph("")
                continue
            for row in table_lines:
                add_paragraph(doc, row, first_indent=False)
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(stripped[2:]))
            i += 1
            continue

        match = numbered_re.match(stripped)
        if match:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(match.group(2)))
            i += 1
            continue

        if table_caption_re.match(stripped):
            add_paragraph(doc, stripped, first_indent=False)
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    doc.save(dst)
    print(dst)


def main() -> None:
    for src_name, dst_name in FILES:
        build_doc(ROOT / src_name, ROOT / dst_name)


if __name__ == "__main__":
    main()
